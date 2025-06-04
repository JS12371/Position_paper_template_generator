import streamlit as st  
import pandas as pd  
from docx import Document  
from docx.oxml import OxmlElement 
from docx.oxml.ns import qn 
from docx.shared import Pt, RGBColor  
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  
import base64  
from io import BytesIO 
import os 
import glob
from docxcompose.composer import Composer  # Import Composer from docxcompose

# Function to convert the DataFrame to Word document  
def mac_num_to_name(mac_num): 
    if mac_num[:2] == '05': 
        return 'WPS Government Health Administrators (J-5)' 
    if mac_num[:2] == '06': 
        return 'National Government Services (J-6)' 
    if mac_num[:2] == '08': 
        return 'WPS Government Health Administrators (J-8)' 
    if mac_num[:2] == '15': 
        return 'CGS Administrators (J-15)' 
    if mac_num[:2] == '01': 
        return 'Noridian Healthcare Solutions c/o Cahaba Safeguard Administrators (J-E)' 
    if mac_num[:2] == '02' or mac_num[:2] == '03': 
        return 'Noridian Healthcare Solutions (J-F)' 
    if mac_num[:2] == '04' or mac_num[:2] == '07': 
        return 'Novitas Solutions, Inc. (J-H)' 
    if mac_num[:2] == '10': 
        return 'Palmetto GBA (J-I)' 
    if mac_num[:2] == '13': 
        return 'National Government Services, Inc. (J-K)' 
    if mac_num[:2] == '12': 
        return 'Novitas Solutions, Inc. (J-L)' 
    if mac_num[:2] == '11': 
        return 'Palmetto GBA c/o National Government Services, Inc. (J-M)' 
    if mac_num[:2] == '09': 
        return 'First Coast Service Options, Inc. (J-N)' 

def format_date(date): 
    year = date[:4] 
    month = date[5:7] 
    day = date[8:10] 
    return f"{month}/{day}/{year}" 

def sanitize_filename(filename):
    invalid_chars = '\\/*?:"<>|'
    for char in invalid_chars:
        filename = filename.replace(char, "")
    filename = filename.replace(" ", "")
    return filename

def get_issue_content(issue, selected_argument):
    issueformatted = sanitize_filename(issue)
    filename = f"IssuestoArgs/{issueformatted}_{selected_argument}.docx"
    if not os.path.exists(filename):
        return None, f"Argument for Issue '{issue}' has not yet been added to the database."
    try:
        doc = Document(filename)
        return doc, None
    except Exception as e:
        return None, f"Error processing issue file: {e}"

def get_possible_arguments(issue):
    issueformatted = sanitize_filename(issue)
    files = glob.glob(f"IssuestoArgs/{issueformatted}_*.docx")
    arguments = [os.path.basename(f).replace(f"{issueformatted}_", "").replace(".docx", "") for f in files]
    return arguments


def extract_exhibits(doc):
    exhibits = []
    in_exhibits_section = False
    for paragraph in doc.paragraphs:
        if "EXHIBITS" in paragraph.text:
            in_exhibits_section = True
        if in_exhibits_section and paragraph.text.startswith("C-"):
            exhibits.append(paragraph)
    return exhibits

def remove_exhibits_from_document(doc):
    in_exhibits_section = False
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        if "EXHIBITS" in paragraph.text:
            in_exhibits_section = True
        if in_exhibits_section:
            paragraphs_to_remove.append(paragraph)
    for paragraph in paragraphs_to_remove:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

def extract_law_regulations(doc):
    law_regulations = {'Law': [], 'Regulations': [], 'Program Instructions': [], 'Other Sources': [], 
                       'United States Statutes': [], 'Judicial Decisions': [], 'Agency Decisions': [], 'SSA': [],
                      'Federal Register': [], 'Agency Instructions': [], 'Case Law': []}
    current_section = None
    in_law_regulations_section = False
    for paragraph in doc.paragraphs:
        if "LAW, REGULATIONS, AND PROGRAM INSTRUCTIONS" in paragraph.text.upper():
            in_law_regulations_section = True
            continue
        if in_law_regulations_section:
            if paragraph.text.startswith("Law"):
                current_section = 'Law'
            elif paragraph.text.startswith("Case Law"):
                current_section = 'Case Law'
            elif paragraph.text.startswith("United States Statutes"):
                current_section = 'United States Statutes'
            elif paragraph.text.startswith("Judicial Decisions"):
                current_section = 'Judicial Decisions'
            elif paragraph.text.startswith("Agency Decisions"):
                current_section = 'Agency Decisions'
            elif paragraph.text.startswith("Federal Register"):
                current_section = 'Federal Register'
            elif paragraph.text.startswith("SSA"):
                current_section = 'SSA'
            elif paragraph.text.startswith("Regulations"):
                current_section = 'Regulations'
            elif paragraph.text.startswith("Program Instructions"):
                current_section = 'Program Instructions'
            elif paragraph.text.startswith("Agency Instructions"):
                current_section = 'Agency Instructions'
            elif paragraph.text.startswith("Other Sources"):
                current_section = 'Other Sources'
            elif paragraph.text.startswith("EXHIBITS"):
                break
            if current_section and paragraph.text.strip() and not paragraph.text.startswith(current_section):
                entries = paragraph.text.split(";")
                for entry in entries:
                    if entry.strip():
                        law_regulations[current_section].append(entry.strip())
    return law_regulations



def remove_law_regulations_from_document(doc):
    in_law_regulations_section = False
    paragraphs_to_remove = []
    for paragraph in doc.paragraphs:
        if "LAW, REGULATIONS, AND PROGRAM INSTRUCTIONS" in paragraph.text.upper():
            in_law_regulations_section = True
        if in_law_regulations_section and not paragraph.text.startswith("EXHIBITS"):
            paragraphs_to_remove.append(paragraph)
        if "EXHIBITS" in paragraph.text.upper():
            in_law_regulations_section = False
    for paragraph in paragraphs_to_remove:
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

def create_word_document(case_data, selected_arguments):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    
    header = doc.add_paragraph('BEFORE THE PROVIDER REIMBURSEMENT REVIEW BOARD')
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = header._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    case_num = case_data['Case Num'].iloc[0] if 'Case Num' in case_data else 'Case Num not found'
    case_name = case_data['Case Name'].iloc[0] if 'Case Name' in case_data else '<input provider name>'
    issue = case_data['Issue'] if 'Issue' in case_data else ['Issue not found']
    cloneissue = [iss for iss in issue]

    tempissue = [i for i in issue]
    issue = tempissue

    transferred_to_case = case_data['Transferred to Case #'] if 'Transferred to Case #' in case_data else ['transferred to case not found']
    temptransferred_to_case = [i for i in transferred_to_case]
    transferred_to_case = temptransferred_to_case

    group_mode = False
    if case_num.endswith('G') or case_num.endswith('C'):
        group_mode = True

    i = 0
    if len(issue) != 1:
        while i < len(issue):
            if transferred_to_case[i] != 'Not in the spreadsheet':
                issue[i] = f"Transferred to case {transferred_to_case[i]}"
            i += 1

    issue = list(dict.fromkeys(issue))
    if issue[0] == 'Issue not found':
        try:
            issue = case_data['Issue Typ'].iloc[0].split(',') if 'Issue Typ' in case_data else ['Issue not found']
        except:
            pass

    provider_numbers = ', '.join(case_data['Provider ID'].unique()) if 'Provider ID' in case_data else 'Provider Numbers not found'
    provider_num_array = case_data['Provider ID'].unique() if 'Provider ID' in case_data else 'Provider Numbers not found'
    if len(provider_num_array) > 1:
        provider_numbers = "Various"
    else:
        pass

    provider_names = ', '.join(case_data['Provider Name'].unique()) if 'Provider Name' in case_data else 'Provider Names not found'
    provider_name_array = case_data['Provider Name'].unique() if 'Provider Name' in case_data else 'Provider Names not found'
    if len(provider_name_array) > 1:
        provider_names = "Various"
    else:
        pass

    case_num = case_data['Case Num'].iloc[0] if 'Case Num' in case_data else 'Case Num not found'
    mac_num = case_data['MAC'].iloc[0] if 'MAC' in case_data else 'MAC not found'
    mac_name = mac_num_to_name(mac_num)

    determination_event_dates = ', '.join([format_date(str(date)[:10]) for date in case_data['Determination Event Date'].unique()]) if 'Determination Event Date' in case_data else 'Determination Event Dates not found'
    det_event_array = case_data['Determination Event Date'].unique() if 'Determination Event Date' in case_data else 'Determination Event Dates not found'
    if len(det_event_array) > 1:
        determination_event_dates = 'Various'
    else:
        pass

    if issue[0].startswith('Transfer'):
        issue.remove(issue[0])
    date_of_appeal = format_date(str(case_data['Appeal Date'].iloc[0])[:10]) if 'Appeal Date' in case_data else 'Date of Appeal not found'
    adj_no = ','.join(case_data['Audit Adj No.'].unique()) if 'Audit Adj No.' in case_data else 'Audit Adj No. not found'
    if 'Group FYE' in case_data:
        year = format_date(case_data['Group FYE'].iloc[0]) if 'Group FYE' in case_data else 'FYE not found'
    else:
        year = format_date(case_data['FYE'].iloc[0]) if 'FYE' in case_data else 'FYE not found'

    table = doc.add_table(rows = 1, cols = 3)

    for cell in table.columns[0].cells:
        cell.width = Pt(260)
    for cell in table.columns[1].cells:
        cell.width = Pt(20)
    for cell in table.columns[2].cells:
        cell.width = Pt(260)

    cell_left = table.cell(0,0)
    cell_left.text = f"\n{case_name}\n\nProvider Numbers: {provider_numbers}\n\n     Provider Names: {provider_names} \n\n vs. \n\n{mac_name}\n     (Medicare Administrative Contractor)\n\n        and \n\n Federal Specialized Services \n     (Appeals Support Contractor)\n"
    run = cell_left.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    cell_middle = table.cell(0,1)
    cell_middle.text = ")\n"*16
    cell_middle.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_middle.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = cell_middle.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    cell_right = table.cell(0,2)
    cell_right.text = f"\n\n\n\n\n\nPRRB Case No. {case_num}\n\nFYE: {year[:]}\n"
    run = cell_right.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    line_para = doc.add_paragraph()
    line_para.add_run()
    p = line_para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '6')
    bottom_bdr.set(qn('w:space'), '1')
    bottom_bdr.set(qn('w:color'), 'auto')
    pBdr.append(bottom_bdr)
    pPr.append(pBdr)

    header = doc.add_paragraph('MEDICARE ADMINISTRATIVE CONTRACTOR’S POSITION PAPER')
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    sub = doc.add_paragraph(f"Submitted by:\n\n<Name>\n{mac_name}\n<Address>\n<Address line 2>")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = sub.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    sub = doc.add_paragraph(f"and\n\n<Reviewer Name>\nFederal Specialized Services, LLC\n1701 S. Racine Avenue\nChicago, IL 60608-4058")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = sub.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()

    table1 = doc.add_table(rows = 1, cols = 2)
    for cell in table1.columns[0].cells:
        cell.width = Pt(500)
    for cell in table1.columns[1].cells:
        cell.width = Pt(40)

    cell_left1 = table1.cell(0,0)
    for cell in table1.columns[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

    cell_left1.text = f"TABLE OF CONTENTS"
    run = cell_left1.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    cell_right1 = table1.cell(0,1)
    cell_right1.text = f"PAGE"
    run = cell_right1.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows = 1, cols = 2)
    for cell in table.columns[0].cells:
        cell.width = Pt(500)
    for cell in table.columns[1].cells:
        cell.width = Pt(40)

    cell_left = table.cell(0,0)
    cell_left.text = f"\nI. INTRODUCTION\n\nII. ISSUES AND ADJUSTMENTS IN DISPUTE\n\nIII. MAC'S POSITION\n\nIV. CITATION OF PROGRAM LAWS, REGULATIONS, INSTRUCTIONS, AND CASES\n\nV. EXHIBITS"
    run = cell_left.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    cell_right = table.cell(0,1)
    cell_right.text = "\n1\n\n2\n\n3\n\n?\n\n\n?"
    run = cell_right.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(f"Case_{case_num}.docx")
    doc.add_page_break()

    header = doc.add_paragraph('I. INTRODUCTION')
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = header.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    header = doc.add_paragraph()
    run = header.add_run()
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.text = f"\nCase Name: {case_name}\n\nProvider Numbers: {provider_numbers}\n\nLead Contractor: {mac_name}\n\nCalendar Year: {year[:]}\n\nPRRB Case Number: {case_num}\n\nDates of Determinations: {determination_event_dates}\n\nDate of Appeal: {date_of_appeal}"

    doc.add_page_break()

    header = doc.add_paragraph('II. ISSUES AND ADJUSTMENTS IN DISPUTE')
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = header.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    header = doc.add_paragraph()
    run = header.add_run()
    run.font.color.rgb = RGBColor(0, 0, 0)

    for i in range(len(issue)):
        if group_mode:
            if len(adj_no) > 1:
                adj_no = "Various"
            header = doc.add_paragraph(f"Issue: {issue[i]}\n\nAdjustment No(s): {adj_no}\n\nApproximate Reimbursement Amount: N/A\n")
            run = header.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            if len(adj_no) > 1:
                adj_no = "Various"
            header = doc.add_paragraph(f"\nIssue {i + 1}: {cloneissue[i]}")
            run = header.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
            if issue[i].startswith("Transferred"):
                header.add_run(f"\n\nDisposition: {issue[i]}")
            header.add_run(f"\n\nAdjustment No(s): {adj_no}\n\nApproximate Reimbursement Amount: N/A\n")

    doc.add_page_break()

    header = doc.add_paragraph("III. MAC'S POSITION")
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = header.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    i = 0
    issue = list(filter(None, issue))

    if issue[0] == 'Issue not found':
        pass
    else:
        all_law_regulations = Document()
        all_exhibits = Document()
        while i < len(issue):
            # Skip issues that are marked as "Transferred"
            if issue[i].startswith("Transferred"):
                i += 1
                continue
            
            header = doc.add_paragraph(f"\nIssue {i+1}: {issue[i]}")
            run = header.add_run()
            run.font.color.rgb = RGBColor(0, 0, 0)
            issue_doc, error = get_issue_content(issue[i], selected_arguments[i])
            if error:
                header = doc.add_paragraph(f"{error}\n")
            else:
                exhibits = extract_exhibits(issue_doc)
                remove_exhibits_from_document(issue_doc)
                law_regulations = extract_law_regulations(issue_doc)
                remove_law_regulations_from_document(issue_doc)
                composer = Composer(doc)
                composer.append(issue_doc)
                for section, paragraphs in law_regulations.items():
                    if paragraphs:
                        header = all_law_regulations.add_paragraph(f"\n{section}:")
                        header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        run = header.runs[0]
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        for paragraph in paragraphs:
                            all_law_regulations.add_paragraph(paragraph)
                if exhibits:
                    header = all_exhibits.add_paragraph(f"\nISSUE: {issue[i]}")
                    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    run = header.runs[0]
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    for exhibit in exhibits:
                        all_exhibits.add_paragraph(exhibit.text)
            run = header.add_run()
            run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1

        # Append all law and regulations to the main document
        if all_law_regulations.paragraphs:
            # Ensure there's no extra blank page before this section
            if doc.paragraphs[-1].text.strip():
                doc.add_page_break()
            header = doc.add_paragraph('IV. LAW, REGULATIONS, AND PROGRAM INSTRUCTIONS')
            header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = header.runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

            composer = Composer(doc)
            composer.append(all_law_regulations)

        # Append all exhibits to the main document at the end
        if all_exhibits.paragraphs:
            header = doc.add_paragraph('\nV. EXHIBITS')
            header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = header.runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

            composer = Composer(doc)
            composer.append(all_exhibits)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                        run.font.color.rgb = RGBColor(0, 0, 0)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
  
def string_processing(s): 
    if pd.isnull(s) or s == '': 
        return "Not in the spreadsheet" 
    return str(s).replace('"', '') 

def find_case_data(df, case_number): 
    case_number = case_number.upper()
    df['Case Num'] = df['Case Num'].map(string_processing) 
    case_data = df[df['Case Num'] == case_number] 
    case_data = case_data.map(string_processing) 
    return case_data 

def get_download_link(file, filename): 
    b64 = base64.b64encode(file).decode() 
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download file</a>' 
    return href 


st.title('QR Position Paper Template Generator')  

# Step 1: Upload Excel file
uploaded_file = st.file_uploader("Choose an Excel file (upload the 061 report)", type=['xlsx', 'xls'])  

# Maintain the loaded DataFrame
if 'df' not in st.session_state:
    st.session_state.df = None

# Define the relevant columns to read
relevant_columns0 = ['Case Num', 'Case Name', 'Issue', 'Transferred to Case #', 'Provider ID', 'Provider Name', 'MAC', 'Determination Event Date', 'Appeal Date', 'Audit Adj No.', 'Issue Typ']
relevant_columns1 = ['Case Num', 'Case Name', 'Issue', 'Transferred to Case #', 'Provider ID', 'Provider Name', 'MAC', 'Determination Event Date', 'Appeal Date', 'Audit Adj No.', 'FYE', 'Issue Typ']
relevant_columns2 = ['Case Num', 'Case Name', 'Issue', 'Transferred to Case #', 'Provider ID', 'Provider Name', 'MAC', 'Determination Event Date', 'Appeal Date', 'Audit Adj No.', 'Group FYE', 'Issue Typ']

# Modify the read_excel call to use only the relevant columns based on the file name
if uploaded_file and st.session_state.df is None:
    # Analyze the file name
    file_name = uploaded_file.name
    
    # Determine which columns to use based on the file name
    if file_name.startswith('045'):
        relevant_columns = relevant_columns2
    elif file_name.startswith('061'):
        relevant_columns = relevant_columns1
    else:
        # Default to relevant_columns1 if no specific condition is met
        relevant_columns = relevant_columns0
    
    try:
        st.session_state.df = pd.read_excel(uploaded_file, usecols=lambda col: col in relevant_columns, engine='calamine')
        if not st.session_state.df.empty:
            st.write('File uploaded successfully')
        else:
            st.write('Failed to read the file.')
    except Exception as e:
        st.write(f'Error reading file: {e}')


# Proceed only if the DataFrame is loaded
if st.session_state.df is not None:
    # Step 2: Enter Case Number
    case_num = st.text_input('Enter Case Number', value=st.session_state.get('case_num', ''))
    find_case_button = st.button('Find Case') 

    # Maintain the loaded case_data
    if 'case_data' not in st.session_state:
        st.session_state.case_data = None

    if case_num and find_case_button:
        st.session_state.case_data = find_case_data(st.session_state.df, case_num)
        st.session_state.selected_arguments = {}
        st.session_state.case_num = case_num

    # Proceed only if the case_data is found
    if st.session_state.case_data is not None:
        if not st.session_state.case_data.empty:
            issues = st.session_state.case_data['Issue'].unique()
            for issue in issues:
                arguments = get_possible_arguments(issue)
                if arguments:
                    selected_argument = st.selectbox(f"Select argument for issue '{issue}'", arguments, key=issue, 
                                                     index=arguments.index(st.session_state.selected_arguments.get(issue, arguments[0])))
                    st.session_state.selected_arguments[issue] = selected_argument
                else:
                    st.session_state.selected_arguments[issue] = ""

            # Step 3: Create Document
            create_doc = st.button('Create Document') 
            if create_doc:
                docx_file = create_word_document(st.session_state.case_data, [st.session_state.selected_arguments[issue] for issue in issues])
                st.markdown(get_download_link(docx_file, f'Case_{case_num}.docx'), unsafe_allow_html=True)
        else:
            st.write('Case not found in the spreadsheet. Please try again with a different case number.')
