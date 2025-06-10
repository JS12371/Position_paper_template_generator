# Note: This code assumes `database.db` exists and contains all relevant case/exhibit/law data.
# This is the integrated version of your Streamlit app with dynamic MAC Section III-V generation.

import streamlit as st
import pandas as pd
import sqlite3
from datetime import datetime
from docx import Document
from io import BytesIO
import base64
import os
import re
import PyPDF2
from docx.oxml import OxmlElement 
from docx.oxml.ns import qn 
from docx.shared import Pt, RGBColor  
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  
from docxcompose.composer import Composer

DB_NAME = "database.db"
import os


def get_all_stratifiers():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT DISTINCT stratifier FROM Cases")
    stratifiers = [row[0] for row in c.fetchall()]
    conn.close()
    return stratifiers

def get_recent_case_data(stratifier, cutoff_date):
    # Get the directory where the current script is located
    current_dir = os.path.dirname(os.path.abspath(__file__))
    facts_path = os.path.join(current_dir, "stratifiers", stratifier, "Facts.txt")
    
    try:
        with open(facts_path, 'r') as f:
            facts = f.read().strip()
    except FileNotFoundError as e:
        print(f"Error: Could not find required files: {e}")
        return []
    
    return [[1, facts, "", ""]]  # Return empty strings for arguments and conclusion

def get_all_laws_and_exhibits_for_stratifier(stratifier):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()

    c.execute("""
        SELECT DISTINCT L.citation
        FROM Cases C
        JOIN CaseLaws CL ON CL.case_id = C.id
        JOIN Laws L ON L.id = CL.law_id
        WHERE C.stratifier = ?
    """, (stratifier,))
    all_laws = [row[0] for row in c.fetchall()]

    c.execute("""
        SELECT DISTINCT E.title
        FROM Cases C
        JOIN CaseExhibits CE ON CE.case_id = C.id
        JOIN Exhibits E ON E.id = CE.exhibit_id
        WHERE C.stratifier = ?
    """, (stratifier,))
    all_exhibits = [row[0] for row in c.fetchall()]

    conn.close()
    return all_laws, all_exhibits


def create_word_document():
    doc = Document()
    doc.add_heading("III. MAC's Position", level=1)
    doc.add_paragraph(st.session_state.section_iii)
    #add a new line
    doc.add_paragraph("")

    doc.add_heading("IV. LAW, REGULATIONS, AND PROGRAM INSTRUCTIONS", level=1)
    for law in st.session_state.section_iv:
        doc.add_paragraph(f"- {law}")

    #add a new line
    doc.add_paragraph("")

    doc.add_heading("V. EXHIBITS", level=1)
    for ex in st.session_state.section_v:
        doc.add_paragraph(f"- {ex}")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def get_download_link(file, filename):
    b64 = base64.b64encode(file).decode()
    return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download Document</a>'

  # Import Composer from docxcompose

# Function to convert the DataFrame to Word document  
def mac_num_to_name(mac_num): 
    if mac_num[:2] == '05': 
        return 'WPS Government Health Administrators (J-5)' 
    if mac_num[:2] == '06': 
        return 'National Government Services (J-6)' 
    if mac_num[:2] == '08': 
        return 'WPS Government Health Administrators (J-8)' 
    if mac_num[:2] == '15': 
        return 'CGS Administrators (J-15)' 
    if mac_num[:2] == '01': 
        return 'Noridian Healthcare Solutions c/o Cahaba Safeguard Administrators (J-E)' 
    if mac_num[:2] == '02' or mac_num[:2] == '03': 
        return 'Noridian Healthcare Solutions (J-F)' 
    if mac_num[:2] == '04' or mac_num[:2] == '07': 
        return 'Novitas Solutions, Inc. (J-H)' 
    if mac_num[:2] == '10': 
        return 'Palmetto GBA (J-I)' 
    if mac_num[:2] == '13': 
        return 'National Government Services, Inc. (J-K)' 
    if mac_num[:2] == '12': 
        return 'Novitas Solutions, Inc. (J-L)' 
    if mac_num[:2] == '11': 
        return 'Palmetto GBA c/o National Government Services, Inc. (J-M)' 
    if mac_num[:2] == '09': 
        return 'First Coast Service Options, Inc. (J-N)' 

def format_date(date): 
    year = date[:4] 
    month = date[5:7] 
    day = date[8:10] 
    return f"{month}/{day}/{year}" 


def extract_issue_input(text, provider_name):
    # Pattern to match text between "Fiscal Year" and "for"
    pattern = r"Fiscal Year\s+(.*?)\s+for"
    
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    
    if match:
        # Clean up the captured text
        issue_input = match.group(1).strip()
        # Replace multiple newlines with a single space
        issue_input = re.sub(r'\n\s*', ' ', issue_input)
        return issue_input
    return None

def create_word_document(case_data, selected_stratifier, provider_name, issue_text=None):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(9)
    
    header = doc.add_paragraph('BEFORE THE PROVIDER REIMBURSEMENT REVIEW BOARD')
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)
    p = header._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

    case_num = case_data['Case Num'].iloc[0] if 'Case Num' in case_data else 'Case Num not found'
    
    # Get the issue from the uploaded PDF if available
    if issue_text:
        # Get the directory where the current script is located
        current_dir = os.path.dirname(os.path.abspath(__file__))
        issue_path = os.path.join(current_dir, "stratifiers", selected_stratifier, "Issue.txt")
        
        try:
            with open(issue_path, 'r') as f:
                issue_template = f.read().strip()
                # Replace <INPUT> with the extracted issue input
                issue = issue_template.replace("<INPUT>", issue_text)
        except FileNotFoundError as e:
            print(f"Error: Could not find required files: {e}")
            issue = "Issue not found"
    else:
        issue = "Issue not found"

    transferred_to_case = case_data['Transferred to Case #'] if 'Transferred to Case #' in case_data else ['transferred to case not found']
    temptransferred_to_case = [i for i in transferred_to_case]
    transferred_to_case = temptransferred_to_case

    group_mode = False
    if case_num.endswith('G') or case_num.endswith('C'):
        group_mode = True

    

    if 'Provider ID' in case_data:
        provider_numbers = ', '.join(case_data['Provider ID'].unique())
        provider_num_array = case_data['Provider ID'].unique()
    elif 'Prov Num' in case_data:
        provider_numbers = ', '.join(case_data['Prov Num'].unique()) 
        provider_num_array = case_data['Prov Num'].unique()
    else:
        provider_numbers = 'Provider Number(s) not found'
        provider_num_array = 'Provider Numbers not found'

    
    if len(provider_num_array) > 1:
        provider_numbers = "Various"
    else:
        pass

    if 'Provider Name' in case_data:
        provider_names = ', '.join(case_data['Provider Name'].unique())
        provider_name_array = case_data['Provider Name'].unique()
    elif 'Firm' in case_data: 
        provider_names = ', '.join(case_data['Firm'].unique())
        provider_name_array = case_data['Firm'].unique()
    else:
        provider_names = 'Provider Name(s) not found'
        provider_name_array = 'Provider Name(s) not found'
    
    if len(provider_name_array) > 1:
        provider_names = "Various"
    else:
        pass

    if 'Est. Reimb. Impact' in case_data:
        reimbursement = 1.0 or 1
        for i in case_data['Est. Reimb. Impact'].unique():
            if float(i) != 1.0 and float(i) != 1:
                reimbursement = i
                break
    else:
        reimbursement = 'N/A'

    case_num = case_data['Case Num'].iloc[0] if 'Case Num' in case_data else 'Case Num not found'
    mac_num = case_data['MAC'].iloc[0] if 'MAC' in case_data else 'MAC not found'
    mac_name = mac_num_to_name(mac_num)

    determination_event_dates = ', '.join([format_date(str(date)[:10]) for date in case_data['Determination Event Date'].unique()]) if 'Determination Event Date' in case_data else 'Determination Event Dates not found'
    det_event_array = case_data['Determination Event Date'].unique() if 'Determination Event Date' in case_data else 'Determination Event Dates not found'
    if len(det_event_array) > 1:
        determination_event_dates = 'Various'
    else:
        pass

    
    if 'Appeal Date' in case_data:
        date_of_appeal = format_date(str(case_data['Appeal Date'].iloc[0])[:10])
    elif 'Appeal Request Date' in case_data:
        date_of_appeal = format_date(str(case_data['Appeal Request Date'].iloc[0])[:10])
    else:
        date_of_appeal = 'Date of Appeal not found'
    
    adj_no = ','.join(case_data['Audit Adj No.'].unique()) if 'Audit Adj No.' in case_data else 'Audit Adj No. not found'
    if 'Group FYE' in case_data:
        year = format_date(case_data['Group FYE'].iloc[0]) if 'Group FYE' in case_data else 'FYE not found'
    else:
        year = format_date(case_data['FYE'].iloc[0]) if 'FYE' in case_data else 'FYE not found'

    table = doc.add_table(rows = 1, cols = 3)

    for cell in table.columns[0].cells:
        cell.width = Pt(260)
    for cell in table.columns[1].cells:
        cell.width = Pt(20)
    for cell in table.columns[2].cells:
        cell.width = Pt(260)

    cell_left = table.cell(0,0)
    cell_left.text = f"\n{provider_name}\n\nProvider Numbers: {provider_numbers}\n\n     (Provider) \n\n vs. \n\n{mac_name}\n     (Medicare Administrative Contractor)\n\n        and \n\n Federal Specialized Services \n     (Appeals Support Contractor)\n"
    run = cell_left.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    cell_middle = table.cell(0,1)
    cell_middle.text = ")\n"*16
    cell_middle.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_middle.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = cell_middle.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    cell_right = table.cell(0,2)
    cell_right.text = f"\n\n\n\n\n\nPRRB Case No. {case_num}\n\nFYE: {year[:]}\n"
    run = cell_right.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    line_para = doc.add_paragraph()
    line_para.add_run()
    p = line_para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_bdr = OxmlElement('w:bottom')
    bottom_bdr.set(qn('w:val'), 'single')
    bottom_bdr.set(qn('w:sz'), '6')
    bottom_bdr.set(qn('w:space'), '1')
    bottom_bdr.set(qn('w:color'), 'auto')
    pBdr.append(bottom_bdr)
    pPr.append(pBdr)

    header = doc.add_paragraph("MEDICARE ADMINISTRATIVE CONTRACTOR'S POSITION PAPER")
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = header.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    sub = doc.add_paragraph(f"Submitted by:\n\n<Name>\n{mac_name}\n<Address>\n<Address line 2>")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = sub.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    sub = doc.add_paragraph(f"and\n\n<Reviewer Name>\nFederal Specialized Services, LLC\n1701 S. Racine Avenue\nChicago, IL 60608-4058")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = sub.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_page_break()

    table1 = doc.add_table(rows = 1, cols = 2)
    for cell in table1.columns[0].cells:
        cell.width = Pt(500)
    for cell in table1.columns[1].cells:
        cell.width = Pt(40)

    cell_left1 = table1.cell(0,0)
    for cell in table1.columns[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)

    cell_left1.text = f"TABLE OF CONTENTS"
    run = cell_left1.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    cell_right1 = table1.cell(0,1)
    cell_right1.text = f"PAGE"
    run = cell_right1.paragraphs[0].runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    table = doc.add_table(rows = 1, cols = 2)
    for cell in table.columns[0].cells:
        cell.width = Pt(500)
    for cell in table.columns[1].cells:
        cell.width = Pt(40)

    cell_left = table.cell(0,0)
    cell_left.text = f"\nI. INTRODUCTION\n\nII. ISSUES AND ADJUSTMENTS IN DISPUTE\n\nIII. MAC'S POSITION\n\nIV. CITATION OF PROGRAM LAWS, REGULATIONS, INSTRUCTIONS, AND CASES\n\nV. EXHIBITS"
    run = cell_left.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    cell_right = table.cell(0,1)
    cell_right.text = "\n1\n\n2\n\n3\n\n?\n\n\n?"
    run = cell_right.paragraphs[0].runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(f"Case_{case_num}.docx")
    doc.add_page_break()

    header = doc.add_paragraph('I. INTRODUCTION')
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = header.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    header = doc.add_paragraph()
    run = header.add_run()
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.text = f"\nProvider Name: {provider_name}\n\nProvider Numbers: {provider_numbers}\n\nLead Contractor: {mac_name}\n\nCalendar Year: {year[:]}\n\nPRRB Case Number: {case_num}\n\nDates of Determinations: {determination_event_dates}\n\nDate of Appeal: {date_of_appeal}"

    doc.add_page_break()

    header = doc.add_paragraph('II. ISSUES AND ADJUSTMENTS IN DISPUTE')
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = header.runs[0]
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    header = doc.add_paragraph()
    run = header.add_run()
    run.font.color.rgb = RGBColor(0, 0, 0)

    if len(adj_no) > 1:
        adj_no = "Various"
    header = doc.add_paragraph(f"Issue: {issue}\n\nAdjustment No(s): {adj_no}\n\nApproximate Reimbursement Amount: {reimbursement}\n")
    run = header.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0) 

    doc.add_page_break()


    doc.add_heading("III. MAC's Position", level=1)
    doc.add_paragraph(st.session_state.section_iii)
    #add a new line
    doc.add_paragraph("")

    doc.add_heading("IV. LAW, REGULATIONS, AND PROGRAM INSTRUCTIONS", level=1)
    for law in st.session_state.section_iv:
        doc.add_paragraph(f"- {law}")

    #add a new line
    doc.add_paragraph("")

    doc.add_heading("V. EXHIBITS", level=1)
    for ex in st.session_state.section_v:
        doc.add_paragraph(f"- {ex}")

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0, 0, 0)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                        run.font.color.rgb = RGBColor(0, 0, 0)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
  
def string_processing(s): 
    if pd.isnull(s) or s == '': 
        return "Not in the spreadsheet" 
    return str(s).replace('"', '') 

def find_case_data(df, case_number): 
    case_number = case_number.upper()
    df['Case Num'] = df['Case Num'].map(string_processing) 
    case_data = df[df['Case Num'] == case_number] 
    case_data = case_data.map(string_processing) 
    return case_data 

def get_download_link(file, filename): 
    b64 = base64.b64encode(file).decode() 
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download file</a>' 
    return href 

def extract_text_from_pdf(pdf_path):
    text = ""
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def check_determination_text(text):
    pattern = r"NOTICE\s+OF\s+QUALITY\s+REPORTING\s+PROGRAM\s+NONCOMPLIANCE\s+DECISION\s+UPHELD"
    return bool(re.search(pattern, text, re.MULTILINE))

def check_iqrp_requirement(text):
    pattern = r"The\s+requirement\s+that\s+facilities\s+participating\s+in\s+the\s+Hospital\s+!QR\s+Program\s+report\s+quality\s+data\s+to\s+CMS\s+is\s+set\s+forth\s+in\s+42\s+Code\s+of\s+Federal\s+Regulations\s+\(CFR\)\s+Part\s+412,\s+Subpart\s+H\."
    pattern2 = r"The\s+requirement\s+that\s+facilities\s+participating\s+in\s+the\s+Hospital\s+IQR\s+Program\s+report\s+quality\s+data\s+to\s+CMS\s+is\s+set\s+forth\s+in\s+42\s+Code\s+of\s+Federal\s+Regulations\s+\(CFR\)\s+Part\s+412,\s+Subpart\s+H\."
    if bool(re.search(pattern, text, re.MULTILINE)):
        return bool(re.search(pattern, text, re.MULTILINE))
    elif bool(re.search(pattern2, text, re.MULTILINE)):
        return bool(re.search(pattern2, text, re.MULTILINE))
    return False


def extract_provider_name(text):
    # Find the determination text first
    determination_pattern = r"NOTICE\s+OF\s+QUALITY\s+REPORTING\s+PROGRAM\s+NONCOMPLIANCE\s+DECISION\s+UPHELD"
    determination_match = re.search(determination_pattern, text, re.MULTILINE)
    
    if determination_match:
        # Get the text after the determination
        text_after_determination = text[determination_match.end():]
        
        # Look for "for" followed by text until "CMS"
        provider_pattern = r"for\s+(.*?)\s+CMS"
        provider_match = re.search(provider_pattern, text_after_determination, re.IGNORECASE | re.DOTALL)
        
        if provider_match:
            return provider_match.group(1).strip()
    return None


st.title('QR Position Paper Template Generator')  

# Add PDF upload section
st.subheader('Upload Final Determination')
uploaded_pdf = st.file_uploader("Choose a PDF file", type=['pdf'])

if uploaded_pdf is not None:
    # Save the uploaded file temporarily
    temp_path = "temp_upload.pdf"
    with open(temp_path, "wb") as f:
        f.write(uploaded_pdf.getvalue())
    
    try:
        # Extract text from PDF
        text = extract_text_from_pdf(temp_path)
        
        # Check if it's a valid determination and IQRP
        is_determination = check_determination_text(text)
        is_iqrp = check_iqrp_requirement(text)
        
        if not is_determination:
            st.error("This is not a valid final determination")
        else:
            if is_iqrp:
                st.success("This is an IQRP determination")
                # Set the stratifier to IQRP
                st.session_state.selected_stratifier = "IQRP"
                # Extract provider name and issue input
            else:
                st.warning("This is not a known determination type")
                
            provider_name = extract_provider_name(text)
            issue_input = extract_issue_input(text, provider_name)
            if issue_input:
                st.session_state.issue_input = issue_input
            
    
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
    
    finally:
        # Clean up the temporary file
        if os.path.exists(temp_path):
            os.remove(temp_path)

# Step 1: Upload Excel file
uploaded_file = st.file_uploader("Choose an Excel file (061 report)", type=['xlsx', 'xls'])  

# Maintain the loaded DataFrame
if 'df' not in st.session_state:
    st.session_state.df = None


relevant_columns1 = ['Case Num', 'Case Name', 'Issue', 'Transferred to Case #', 'Provider ID', 'Prov Num', 'Provider Name', 'Firm', 'MAC', 'Determination Event Date', 'Appeal Date', 'Appeal Request Date', 'Audit Adj No.', 'FYE', 'Issue Typ', 'Est. Reimb. Impact']
 
# Modify the read_excel call to use only the relevant columns based on the file name
if uploaded_file and st.session_state.df is None:
    # Analyze the file name
    file_name = uploaded_file.name
    
    # Determine which columns to use based on the file name
    relevant_columns = relevant_columns1
    if file_name.startswith('045'):
        st.write("Please upload a 061 report")
    elif file_name.startswith('061'):
        relevant_columns = relevant_columns1
    else:
        # Default to relevant_columns1 if no specific condition is met
        relevant_columns = relevant_columns1
    
    try:
        st.session_state.df = pd.read_excel(uploaded_file, usecols=lambda col: col in relevant_columns, engine='calamine')
        if not st.session_state.df.empty:
            st.write('File uploaded successfully')
        else:
            st.write('Failed to read the file.')
    except Exception as e:
        st.write(f'Error reading file: {e}')


# Proceed only if the DataFrame is loaded
if st.session_state.df is not None:
    # Step 2: Enter Case Number
    case_num = st.text_input('Enter Case Number', value=st.session_state.get('case_num', ''))
    find_case_button = st.button('Find Case') 

    # Maintain the loaded case_data
    if 'case_data' not in st.session_state:
        st.session_state.case_data = None

    if case_num and find_case_button:
        st.session_state.case_data = find_case_data(st.session_state.df, case_num)
        st.session_state.selected_arguments = {}
        st.session_state.case_num = case_num

    # Proceed only if the case_data is found
    if st.session_state.case_data is not None:
        if not st.session_state.case_data.empty:
            cutoff_date = st.date_input("Select cutoff date [date of the appeal request] (papers on or after this date will be used)").strftime("%Y-%m-%d")
            stratifier_options = get_all_stratifiers()
            
            # Use the selected stratifier from PDF upload if available
            default_stratifier = st.session_state.get('selected_stratifier', stratifier_options[0] if stratifier_options else None)
            if is_iqrp:
                selected_stratifier = "IQRP"
            else:
                selected_stratifier = "notyetsupported"
            if st.button("Generate MAC Position Sections"):
                case_rows = get_recent_case_data(selected_stratifier, cutoff_date)
                if not case_rows:
                    st.warning("No eligible cases found for this stratifier and date.")
                else:
                    facts_gen, args_gen, concl_gen = case_rows[0][1:]

                    laws, exhibits = get_all_laws_and_exhibits_for_stratifier(selected_stratifier)

                    st.session_state.section_iii = f"A. Facts:\n{facts_gen}\n\nB. Arguments:\n\nC. Conclusion:"
                    st.session_state.section_iv = laws
                    st.session_state.section_v = exhibits

                    st.success("Sections III-V generated and ready.")

            if "section_iii" in st.session_state and st.button("Download Word Document"):
                issue_input = st.session_state.get('issue_input')
                file_data = create_word_document(case_data = st.session_state.case_data, issue_text = issue_input, selected_stratifier = selected_stratifier, provider_name = provider_name)
                st.markdown(get_download_link(file_data, "MAC_Position_Paper.docx"), unsafe_allow_html=True)
        else:
            st.write('Case not found in the spreadsheet. Please try again with a different case number.')






